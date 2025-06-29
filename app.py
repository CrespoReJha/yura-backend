from flask import Flask, request, jsonify
import whisper
import os
from werkzeug.utils import secure_filename
from pydantic import BaseModel, Field
from typing import List, Optional
from enum import Enum
from google import genai
from dotenv import load_dotenv
from docx import Document
import json
from datetime import datetime
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
import io

# --- Configuración inicial ---
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")

app = Flask(__name__)
model_whisper = whisper.load_model("base")
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# --- Esquemas Pydantic ---
class TipoInforme(str, Enum):
    EVOLUCION_GERIATRICA = "EVOLUCION_GERIATRICA"

class Paciente(BaseModel):
    apellido_paterno: str = Field(..., description="Apellido paterno del paciente")
    apellido_materno: str = Field(..., description="Apellido materno del paciente")
    nombres: str = Field(..., description="Nombres del paciente")
    edad: int = Field(..., description="Edad en años")
    sexo: str = Field(..., description="Sexo del paciente (masculino/femenino)")
    n_cama: int = Field(..., description="Número de cama asignada")
    n_historia: int = Field(..., description="Número de historia clínica")

class SignosVitales(BaseModel):
    PA: str = Field(..., description="Presión arterial en mmHg (ej: 120/80)")
    FC: int = Field(..., description="Frecuencia cardíaca en latidos por minuto")
    FR: int = Field(..., description="Frecuencia respiratoria en respiraciones por minuto")
    O2: int = Field(..., description="Porcentaje de saturación de oxígeno")

class Neurologico(BaseModel):
    estado: str = Field(..., description="Estado de conciencia y respuesta")
    glasgow: str = Field(..., description="Puntaje de Glasgow (ej: 15 sobre 15 -> 15/15)")
    foco_motor: str = Field(..., description="Presencia de foco motor o sensitivo")

class Evolucion(BaseModel):
    estado_general: str = Field(..., description="Resumen del estado general del paciente")
    EFG: str = Field(..., description="Exploración física general")
    EFR: Optional[str] = Field("", description="Exploración física regional (opcional)")
    cuello: str
    torax_anterior: str
    torax_posterior: str
    abdomen: str
    genitourinario: str
    extremidades: str
    neurologico: Neurologico

class Ingresos(BaseModel):
    VO: int
    VP: int
    AM: int
    OTROS: int
    TOTAL: int

class Egresos(BaseModel):
    D: int
    C: int
    PI: int
    OTROS: int
    TOTAL: int

class InformeGeriatrico(BaseModel):
    tipo_informe: TipoInforme = Field(..., description="Tipo de informe clínico")
    paciente: Paciente
    fecha_hora: str = Field(..., description="Fecha y hora del informe en formato ISO 8601")
    signos_vitales: SignosVitales
    diagnosticos: List[str]
    evolucion: Evolucion
    ingresos: Ingresos
    egresos: Egresos
    ordenes_medicas: List[str]
    BH: str = Field(..., description="Resultados de bioquímica hematológica")
    RD: str = Field(..., description="Resultados de rayos X o diagnóstico por imágenes")
    descripcion_paciente: str = Field(..., description="Descripción clínica del paciente usando sexo, edad y días de internación")

def traducir_dia(dia_en):
    """Traduce nombres de días del inglés al español"""
    dias = {
        "Monday": "Lunes", "Tuesday": "Martes", "Wednesday": "Miércoles",
        "Thursday": "Jueves", "Friday": "Viernes", "Saturday": "Sábado", "Sunday": "Domingo"
    }
    return dias.get(dia_en, dia_en)

def agregar_texto_negrita(paragraph, texto, fuente=10):
    """Agrega texto en negrita a un párrafo"""
    run = paragraph.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(fuente)

def agregar_texto(paragraph, texto, fuente=10):
    """Agrega texto normal a un párrafo"""
    run = paragraph.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(fuente)

def generar_docx_desde_json(data: dict, nombre_archivo: str = "informe.docx"):
    """Genera un documento Word desde datos JSON estructurados"""
    doc = Document()
    
    # Configurar márgenes de 1cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.39)     # 1cm = 0.39 inches
        section.bottom_margin = Inches(0.39)
        section.left_margin = Inches(0.39)
        section.right_margin = Inches(0.39)

    # -------------------------
    # Tabla: Datos del paciente
    # -------------------------
    encabezado_paciente = ["AP PATERNO", "AP MATERNO", "NOMBRES", "HISTORIA CLÍNICA", "NO CAMA"]
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Configurar anchos proporcionales para la tabla de paciente
    table.autofit = False
    ancho_total_paciente = Inches(7.28)  # Ancho disponible con márgenes de 1cm
    anchos_paciente = [Inches(1.5), Inches(1.5), Inches(1.8), Inches(1.24), Inches(1.24)]
    
    for i, ancho in enumerate(anchos_paciente):
        table.columns[i].width = ancho
    
    hdr_cells = table.rows[0].cells
    for i, titulo in enumerate(encabezado_paciente):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        agregar_texto_negrita(p, titulo, 9)

    paciente = data.get('paciente', {})
    row_cells = table.add_row().cells
    datos_paciente = [
        paciente.get('apellido_paterno', ''),
        paciente.get('apellido_materno', ''),
        paciente.get('nombres', ''),
        str(paciente.get('n_historia', '')),
        str(paciente.get('n_cama', ''))
    ]
    
    for i, dato in enumerate(datos_paciente):
        p = row_cells[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        agregar_texto(p, dato, 9)

    doc.add_paragraph()  # espacio

    # -------------------------
    # Tabla principal: Fecha | Notas | Ordenes
    # -------------------------
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = 'Table Grid'
    tabla.autofit = False

    # Ajustar anchos para que no se pasen de los bordes
    ancho_total = Inches(7.28)  # Ancho disponible con márgenes de 1cm
    tabla.columns[0].width = Inches(1.8)   # FECHA Y HORA
    tabla.columns[1].width = Inches(3.8)   # NOTAS DE EVOLUCIÓN
    tabla.columns[2].width = Inches(1.68)  # ÓRDENES MÉDICAS

    encabezados = ["FECHA Y HORA", "NOTAS DE EVOLUCIÓN", "ÓRDENES MÉDICAS"]
    for i, texto in enumerate(encabezados):
        p = tabla.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        agregar_texto_negrita(p, texto, 9)

    # Agregar fila de contenido
    fila_contenido = tabla.add_row()

    # -------------------------
    # Columna 1: Fecha y Signos
    # -------------------------
    fecha_iso = data.get('fecha_hora', '')
    try:
        dt = datetime.fromisoformat(fecha_iso)
        dia_semana = traducir_dia(dt.strftime("%A"))
        fecha_formateada = dt.strftime("%d/%m/%y")
        hora = dt.strftime("%H:%M")
    except Exception:
        dia_semana = ""
        fecha_formateada = ""
        hora = ""

    signos_vitales = data.get('signos_vitales', {})
    cell_fecha = fila_contenido.cells[0]
    p = cell_fecha.paragraphs[0]

    if dia_semana:
        agregar_texto_negrita(p, f"{dia_semana}\n", 8)
    if fecha_formateada:
        agregar_texto(p, f"{fecha_formateada}\n", 8)
    if hora:
        agregar_texto(p, f"{hora}\n\n", 8)
    
    for key in ['PA', 'FC', 'FR', 'O2']:
        valor = signos_vitales.get(key)
        if valor:
            agregar_texto_negrita(p, f"{key}: ", 8)
            agregar_texto(p, f"{valor}\n", 8)

    # -------------------------
    # Columna 2: Notas de Evolución
    # -------------------------
    descripcion = data.get("descripcion_paciente", "")
    diagnosticos = data.get("diagnosticos", [])
    evolucion = data.get('evolucion', {})
    cell_notas = fila_contenido.cells[1]
    p_notas = cell_notas.paragraphs[0]

    if descripcion:
        agregar_texto(p_notas, descripcion + "\n\n", 8)

    if diagnosticos:
        agregar_texto_negrita(p_notas, "DIAGNÓSTICOS:\n", 8)
        for diag in diagnosticos:
            agregar_texto(p_notas, f"• {diag.upper()}\n", 8)
        agregar_texto(p_notas, "\n", 8)

    if evolucion.get("estado_general"):
        agregar_texto_negrita(p_notas, "S: ", 8)
        agregar_texto(p_notas, f"{evolucion['estado_general']}\n\n", 8)

    # Campos de evolución física
    campos_fisicos = {
        "cuello": "Cuello",
        "torax_anterior": "Tórax anterior",
        "torax_posterior": "Tórax posterior", 
        "abdomen": "Abdomen",
        "genitourinario": "Genitourinario",
        "extremidades": "Extremidades"
    }
    
    for key, label in campos_fisicos.items():
        valor = evolucion.get(key)
        if valor:
            agregar_texto_negrita(p_notas, f"{label}: ", 8)
            agregar_texto(p_notas, f"{valor}\n", 8)

    if evolucion.get("EFG"):
        agregar_texto_negrita(p_notas, "\nEFG: ", 8)
        agregar_texto(p_notas, f"{evolucion['EFG']}\n", 8)

    # Examen neurológico
    neurologico = evolucion.get("neurologico", {})
    if neurologico and any(neurologico.values()):
        enb_texto = []

        if neurologico.get("estado"):
            enb_texto.append(neurologico["estado"])
        if neurologico.get("foco_motor"):
            enb_texto.append(neurologico["foco_motor"])
        if neurologico.get("glasgow"):
            enb_texto.append(f"Glasgow {neurologico['glasgow']}")

        if enb_texto:
            agregar_texto_negrita(p_notas, "\nENB: ", 8)
            agregar_texto(p_notas, ", ".join(enb_texto) + "\n", 8)

    # Resultados de laboratorio e imágenes
    bh = data.get("BH", "")
    rd = data.get("RD", "")
    if bh:
        agregar_texto_negrita(p_notas, "\nBH: ", 8)
        agregar_texto(p_notas, f"{bh}\n", 8)
    if rd:
        agregar_texto_negrita(p_notas, "\nRD: ", 8)
        agregar_texto(p_notas, f"{rd}\n", 8)

    # -------------------------
    # Columna 3: Órdenes Médicas
    # -------------------------
    ordenes = data.get("ordenes_medicas", [])
    cell_ordenes = fila_contenido.cells[2]
    p_ordenes = cell_ordenes.paragraphs[0]
    
    for i, orden in enumerate(ordenes):
        # Dividir órdenes largas en múltiples líneas si es necesario
        orden_text = f"{i + 1}. {orden}"
        agregar_texto(p_ordenes, orden_text + "\n", 8)

    # Ingresos y Egresos
    ingresos = data.get("ingresos", {})
    egresos = data.get("egresos", {})
    
    if any(ingresos.values()) or any(egresos.values()):
        agregar_texto(p_ordenes, "\n", 8)
        
        if any(ingresos.values()):
            agregar_texto_negrita(p_ordenes, "INGRESOS:\n", 8)
            for key, val in ingresos.items():
                if val:
                    agregar_texto(p_ordenes, f"{key}: {val}\n", 8)
        
        if any(egresos.values()):
            agregar_texto_negrita(p_ordenes, "\nEGRESOS:\n", 8)
            for key, val in egresos.items():
                if val:
                    agregar_texto(p_ordenes, f"{key}: {val}\n", 8)

    # -------------------------
    # Guardar documento
    # -------------------------
    doc.save(nombre_archivo)
    print(f"✅ Documento guardado como: {nombre_archivo}")

# --- Endpoint principal ---
@app.route("/procesar_informe", methods=["POST"])
def procesar_informe():
    """Procesa un archivo de audio y genera un informe médico estructurado"""
    if 'audio' not in request.files or 'tipo_informe' not in request.form:
        return jsonify({"error": "Faltan datos: audio o tipo_informe"}), 400

    tipo_informe = request.form['tipo_informe']
    try:
        tipo_enum = TipoInforme(tipo_informe)
    except ValueError:
        return jsonify({"error": f"Tipo de informe inválido: {tipo_informe}"}), 400

    audio = request.files['audio']
    filename = secure_filename(audio.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    audio.save(filepath)

    try:
        result = model_whisper.transcribe(filepath)
        texto_crudo = result["text"]
        os.remove(filepath)
    except Exception as e:
        return jsonify({"error": f"Error al transcribir audio: {str(e)}"}), 500

    prompt = f"""
Quiero que estructures este texto clínico en formato JSON. El tipo de informe es "{tipo_enum.value}".

Usa la siguiente lógica:
- Extrae los datos del paciente y colócalos en el campo "paciente".
- Organiza signos vitales en el campo "signos_vitales".
- Coloca los diagnósticos como una lista en "diagnosticos".
- Llena "evolucion", "ingresos", "egresos" y "ordenes_medicas" según el contenido médico.
- Usa el formato de fecha ISO 8601 para el campo "fecha_hora".
- No inventes información: si algún campo no está presente, devuélvelo vacío (ej: "" o 0).
- A partir del sexo, edad y día de internación del paciente, construye un campo llamado "descripcion_paciente" con esta estructura:
  "Paciente de sexo {{sexo}} de {{edad}} años de edad en su {{dia_internacion}} día de internación con los diagnósticos de:"

Texto del informe:
\"\"\"{texto_crudo}\"\"\"

Genera una salida JSON estructurada según el esquema del tipo de informe.
"""

    try:
        client = genai.Client(api_key=API_KEY)
        response = client.models.generate_content(
            model="gemini-2.5-pro",
            contents=prompt,
            config={
                "response_mime_type": "application/json",
                "response_schema": InformeGeriatrico
            }
        )
        data_json = json.loads(response.text)
        file_path = "informe.docx"
        generar_docx_desde_json(data_json, nombre_archivo=file_path)

        with open(file_path, "rb") as f:
            doc_bytes = f.read()
            doc_base64 = base64.b64encode(doc_bytes).decode("utf-8")

        return jsonify({
            "json_generado": data_json,
            "documento_base64": doc_base64
        })


        return response.text
    except Exception as e:
        return jsonify({"error": f"Error al generar respuesta con Gemini: {str(e)}"}), 500

# --- Iniciar servidor ---
if __name__ == "__main__":
    app.run(debug=True)